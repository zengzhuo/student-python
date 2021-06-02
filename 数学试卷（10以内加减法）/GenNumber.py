from docx import Document
from docx.shared import Inches
from docx.shared import Cm
import random
from datetime import datetime

document = Document()
document.add_heading('数学试卷(10以内加减法)', 0)
document.add_paragraph(
    '            班级：__________  姓名：_________     得分：___________')
print("准备生成试卷")
# 定义算式集合
formulaSet = set()
# 返回一个加减法算式


def formula():
    flag = False
    symbolType = random.randint(1, 2)
    while flag == False:
        num1 = random.randint(1, 10)
        num2 = random.randint(1, 10)
        if(num1 == num2):
            num2 = random.randint(1, 10)

        if(symbolType == 1):
            if(num1 + num2 <= 10):
                flag = True
        else:
            if(num1 - num2 > 0):
                flag = True
    # 定义运算符类型 1 = + ， 2 = - ，3 = *

    if(symbolType == 1):
        return str(num1) + " + " + str(num2) + " = "
    else:
        return str(num1) + " - " + str(num2) + " = "


i = 1
while i <= 50:
    flag = False
    while flag == False:
        symbol = formula()
        print("生成算式：" + symbol)
        # print(formulaSet)
        flag = symbol in formulaSet
        print("算式是否存在：" + str(flag))
        if(flag == False):
            print("开始生成第*****" + str(i) + "个算式:" + symbol)
            formulaSet.add(symbol)
            flag = True
        else:
            flag = False

    i = i + 1
print("测试的内容总数：" + str(len(formulaSet)))
print(formulaSet)
print("***********开始有序写入算式到word文档中的表格进行填充***********")
table = document.add_table(rows=13, cols=4, style='Table Grid')
# table.autofit = True
row = 0
col = 0
colCount = 4
for symbol in iter(formulaSet):
    if(col >= 4):
        row = row + 1
        col = 0
    table.cell(row, col).text = symbol
    table.cell(row, col).height = Cm(3000)
    col = col + 1
for row in table.rows:
    row.height = Cm(0.1)
dt = datetime.now()
document.save("数学试卷" + dt.strftime('%Y%m%d%H%M%S') + ".docx")
